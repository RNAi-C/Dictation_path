"""
PathDictate v0.2.1 — Pathology Dictation Assistant
Local · Offline · Privacy-first · Real-time transcription
"""

import sys
import os
import re
import json
import threading
import queue
import time
import datetime
import numpy as np
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).parent))

import sounddevice as sd
from loguru import logger

from config import PathologyDictationConfig
from audio_recorder import AudioRecorder
from transcriber import PathologyTranscriber
from terminology_corrector import TerminologyCorrector
from clipboard_handler import ClipboardHandler
from rewriter import LocalRewriter, scan_models
from ollama_client import OllamaClient, OllamaError, OllamaConnectionError
from rewrite_service import RewriteService

APP_VERSION = "0.2.1"

# ── Colour palette ─────────────────────────────────────────────────────────────
BG        = "#F4F6FB"
BG2       = "#FFFFFF"
BG3       = "#EAECF5"
BORDER    = "#CDD0E3"

ACCENT    = "#1A5FB4"
ACCENT_H  = "#124A8C"
ACCENT_L  = "#D6E4F7"

RED_REC   = "#C0392B"
RED_H     = "#A93226"
RED_SOFT  = "#FDECEA"

GREEN_OK   = "#1A7A4A"
GREEN_SOFT = "#E8F5EE"

AMBER      = "#956E00"
AMBER_SOFT = "#FFF8E1"

TEXT      = "#1A1C2E"
TEXT_MED  = "#3D3F58"
TEXT_DIM  = "#6B6D88"
TEXT_LIVE = "#1A5FB4"

# ── Tactile button palette ─────────────────────────────────────────────────────
BTN_BG  = "#DEE1EE"
BTN_HL  = "#CDD1E6"
BTN_PR  = "#BBC0D8"
BTN_FG  = TEXT

REC_BG  = "#C0392B"
REC_HL  = "#A93226"
REC_PR  = "#922B21"

SAV_BG  = "#1A7A4A"
SAV_HL  = "#15603A"
SAV_PR  = "#10482C"

OPN_BG  = "#2D5A9E"
OPN_HL  = "#234880"
OPN_PR  = "#1A3762"

RWR_BG  = "#1A5FB4"
RWR_HL  = "#124A8C"
RWR_PR  = "#0E3A6B"

# ── Typography ────────────────────────────────────────────────────────────────
FONT_UI    = ("Segoe UI", 11)
FONT_BIG   = ("Segoe UI", 13, "bold")
FONT_HEAD  = ("Segoe UI", 16, "bold")
FONT_SMALL = ("Segoe UI", 10)
FONT_MONO  = ("Consolas", 13)

FONT_SIZE_DEFAULT = 13
FONT_SIZE_MIN     = 9
FONT_SIZE_MAX     = 26

STREAM_CHUNK_SEC = 3.0
STREAM_BEAM      = 2


# ── Tactile button factory ─────────────────────────────────────────────────────

def mk_btn(parent, text, cmd, style="default", size="normal", **kw):
    """
    Tactile tk.Button with hover / press visual feedback.
    style : "default" | "record" | "save" | "open" | "rewrite" | "danger" | "subtle"
    size  : "small" | "normal" | "large"
    """
    palettes = {
        "default": (BTN_BG, BTN_FG,   BTN_HL,  BTN_PR),
        "record":  (REC_BG, "white",  REC_HL,  REC_PR),
        "save":    (SAV_BG, "white",  SAV_HL,  SAV_PR),
        "open":    (OPN_BG, "white",  OPN_HL,  OPN_PR),
        "rewrite": (RWR_BG, "white",  RWR_HL,  RWR_PR),
        "danger":  (REC_BG, "white",  REC_HL,  REC_PR),
        "subtle":  (BG2,    TEXT_MED, ACCENT_L, BORDER),
    }
    bg, fg, bg_h, bg_p = palettes.get(style, palettes["default"])

    fs_map = {"small": 9, "normal": 10, "large": 13}
    fs     = fs_map.get(size, 10)
    bold   = style in ("record", "save", "open", "rewrite")
    font   = ("Segoe UI", fs, "bold") if bold else ("Segoe UI", fs)

    padx = kw.pop("padx", {"small": 8, "normal": 12, "large": 22}.get(size, 12))
    pady = kw.pop("pady", {"small": 4, "normal": 7,  "large": 14}.get(size, 7))
    bd   = kw.pop("bd", 3 if style == "record" else 2)

    btn = tk.Button(
        parent, text=text, font=font,
        bg=bg, fg=fg,
        activebackground=bg_h, activeforeground=fg,
        relief="raised", bd=bd,
        padx=padx, pady=pady,
        cursor="hand2", command=cmd, **kw)

    def _enter(_):
        if str(btn["state"]) != "disabled":
            btn.config(bg=bg_h)
    def _leave(_):
        if str(btn["state"]) != "disabled":
            btn.config(bg=bg)
    def _press(_):
        if str(btn["state"]) != "disabled":
            btn.config(bg=bg_p, relief="sunken")
    def _release(_):
        if str(btn["state"]) != "disabled":
            btn.config(bg=bg_h, relief="raised")

    btn.bind("<Enter>",           _enter)
    btn.bind("<Leave>",           _leave)
    btn.bind("<ButtonPress-1>",   _press)
    btn.bind("<ButtonRelease-1>", _release)
    return btn


# ── VU meter ──────────────────────────────────────────────────────────────────

class VUMeter(tk.Canvas):
    def __init__(self, parent, **kw):
        super().__init__(parent, height=14, bg=BG3,
                         highlightthickness=1, highlightbackground=BORDER, **kw)
        self._bar  = self.create_rectangle(0, 2, 0, 12, fill=GREEN_OK, outline="")
        self._peak = 0.0
        self._t    = 0.0
        self.bind("<Configure>", lambda _: self._draw(0.0))

    def update(self, rms: float):
        level = min(rms * 7, 1.0)
        now   = time.monotonic()
        if level > self._peak:
            self._peak, self._t = level, now
        elif now - self._t > 0.4:
            self._peak = max(0.0, self._peak - 0.06)
        self._draw(level)

    def reset(self):
        self._peak = 0.0
        self._draw(0.0)

    def _draw(self, level):
        w = self.winfo_width() or 300
        x = int(w * level)
        c = GREEN_OK if level < 0.65 else (AMBER if level < 0.88 else RED_REC)
        self.coords(self._bar, 0, 2, x, 12)
        self.itemconfig(self._bar, fill=c)


# ── Terminology editor ────────────────────────────────────────────────────────

class TerminologyEditor(tk.Toplevel):
    def __init__(self, parent, dict_path: Path, on_save):
        super().__init__(parent)
        self.title("Terminology Editor")
        self.configure(bg=BG)
        self.geometry("720x560")
        self.minsize(580, 420)
        self.resizable(True, True)
        self.transient(parent)
        self.grab_set()
        self._path    = dict_path
        self._on_save = on_save
        self._data    = {}
        self._load()
        self._build()
        self._populate()

    def _load(self):
        try:
            with open(self._path, encoding="utf-8") as f:
                self._data = json.load(f)
        except Exception:
            self._data = {}

    def _build(self):
        root = ttk.Frame(self, padding=16)
        root.pack(fill="both", expand=True)
        root.columnconfigure(0, weight=1)
        root.rowconfigure(2, weight=1)

        tk.Label(root, text="Terminology Corrections",
                 font=("Segoe UI", 13, "bold"), bg=BG, fg=ACCENT)\
            .grid(row=0, column=0, sticky="w")
        tk.Label(root,
                 text="Double-click a row to edit.  Press Delete to remove selected.",
                 font=FONT_SMALL, bg=BG, fg=TEXT_DIM)\
            .grid(row=1, column=0, sticky="w", pady=(2, 10))

        tf = tk.Frame(root, bg=BORDER, bd=1)
        tf.grid(row=2, column=0, sticky="nsew")
        tf.columnconfigure(0, weight=1)
        tf.rowconfigure(0, weight=1)

        self._tree = ttk.Treeview(tf, columns=("spoken", "corrected"),
                                   show="headings", selectmode="browse")
        self._tree.heading("spoken",    text="You say  (speech-to-text hears)")
        self._tree.heading("corrected", text="Corrected to")
        self._tree.column("spoken",    width=320, anchor="w", minwidth=180)
        self._tree.column("corrected", width=290, anchor="w", minwidth=160)
        vsb = ttk.Scrollbar(tf, orient="vertical", command=self._tree.yview)
        self._tree.configure(yscrollcommand=vsb.set)
        self._tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")

        s = ttk.Style(self)
        s.configure("Edit.Treeview",
                     background=BG2, foreground=TEXT,
                     fieldbackground=BG2, rowheight=28, font=FONT_UI, borderwidth=0)
        s.configure("Edit.Treeview.Heading",
                     background=ACCENT_L, foreground=ACCENT,
                     font=("Segoe UI", 10, "bold"), relief="flat", padding=6)
        s.map("Edit.Treeview",
              background=[("selected", ACCENT)],
              foreground=[("selected", "white")])
        self._tree.configure(style="Edit.Treeview")
        self._tree.bind("<Double-1>", self._on_double_click)
        self._tree.bind("<Delete>",   lambda _: self._delete_row())

        add = tk.Frame(root, bg=BG)
        add.grid(row=3, column=0, sticky="ew", pady=(12, 0))
        add.columnconfigure(1, weight=1)
        add.columnconfigure(4, weight=1)

        def lbl(t):
            return tk.Label(add, text=t, font=FONT_SMALL, bg=BG, fg=TEXT_MED)

        lbl("You say:").grid(row=0, column=0, padx=(0, 6))
        self._entry_spoken = tk.Entry(
            add, bg=BG2, fg=TEXT, insertbackground=TEXT,
            relief="solid", bd=1, font=FONT_UI, highlightthickness=0)
        self._entry_spoken.grid(row=0, column=1, sticky="ew", ipady=5)
        tk.Label(add, text="→", font=("Segoe UI", 12, "bold"),
                 bg=BG, fg=ACCENT).grid(row=0, column=2, padx=12)
        lbl("Correct as:").grid(row=0, column=3, padx=(0, 6))
        self._entry_corrected = tk.Entry(
            add, bg=BG2, fg=TEXT, insertbackground=TEXT,
            relief="solid", bd=1, font=FONT_UI, highlightthickness=0)
        self._entry_corrected.grid(row=0, column=4, sticky="ew", ipady=5)
        mk_btn(add, "Add / Update", self._add_or_update,
               style="open", size="small").grid(row=0, column=5, padx=(12, 0))
        self._entry_spoken.bind("<Return>",    lambda _: self._entry_corrected.focus())
        self._entry_corrected.bind("<Return>", lambda _: self._add_or_update())

        bot = tk.Frame(root, bg=BG)
        bot.grid(row=4, column=0, sticky="ew", pady=(12, 0))
        bot.columnconfigure(1, weight=1)
        mk_btn(bot, "🗑  Delete Selected", self._delete_row,
               style="danger", size="small").grid(row=0, column=0)
        self._count_lbl = tk.Label(bot, text="", font=FONT_SMALL,
                                    bg=BG, fg=TEXT_DIM)
        self._count_lbl.grid(row=0, column=1, padx=14, sticky="w")
        mk_btn(bot, "✓  Save & Close", self._save_close,
               style="save", size="small").grid(row=0, column=2)
        mk_btn(bot, "Cancel", self.destroy,
               style="subtle", size="small").grid(row=0, column=3, padx=(8, 0))

    def _populate(self):
        for row in self._tree.get_children():
            self._tree.delete(row)
        for spoken, corrected in sorted(self._data.items()):
            self._tree.insert("", "end", values=(spoken, corrected))
        self._count_lbl.config(
            text=f"{len(self._tree.get_children())} entries")

    def _add_or_update(self):
        spoken    = self._entry_spoken.get().strip().lower()
        corrected = self._entry_corrected.get().strip()
        if not spoken or not corrected:
            messagebox.showwarning("Empty field",
                                   "Both fields are required.", parent=self)
            return
        self._data[spoken] = corrected
        self._populate()
        for iid in self._tree.get_children():
            if self._tree.item(iid, "values")[0] == spoken:
                self._tree.selection_set(iid)
                self._tree.see(iid)
                break
        self._entry_spoken.delete(0, "end")
        self._entry_corrected.delete(0, "end")
        self._entry_spoken.focus()

    def _delete_row(self):
        sel = self._tree.selection()
        if not sel:
            return
        spoken = self._tree.item(sel[0], "values")[0]
        if messagebox.askyesno("Delete", f"Remove '{spoken}'?", parent=self):
            self._data.pop(spoken, None)
            self._populate()

    def _on_double_click(self, _=None):
        sel = self._tree.selection()
        if not sel:
            return
        spoken, corrected = self._tree.item(sel[0], "values")
        self._entry_spoken.delete(0, "end")
        self._entry_spoken.insert(0, spoken)
        self._entry_corrected.delete(0, "end")
        self._entry_corrected.insert(0, corrected)
        self._entry_corrected.focus()

    def _save_close(self):
        try:
            with open(self._path, "w", encoding="utf-8") as f:
                json.dump(self._data, f, indent=2, ensure_ascii=False)
            self._on_save()
            self.destroy()
        except Exception as e:
            messagebox.showerror("Save failed", str(e), parent=self)


# ── Voice command processor ───────────────────────────────────────────────────

class VoiceCommandProcessor:
    DEFAULT_COMMANDS: dict = {
        "enter":         "\n",
        "new line":      "\n",
        "new paragraph": "\n\n",
        "finish case":   "\n**************\n",
    }

    def __init__(self, config_path: Optional[Path] = None):
        self.commands    = dict(self.DEFAULT_COMMANDS)
        self.config_path = config_path
        if config_path and config_path.exists():
            self._load(config_path)
        self._pattern = self._compile()

    def reload(self):
        self.commands = dict(self.DEFAULT_COMMANDS)
        if self.config_path and self.config_path.exists():
            self._load(self.config_path)
        self._pattern = self._compile()

    def _load(self, path: Path):
        try:
            with open(path, encoding="utf-8") as f:
                extra = json.load(f)
            self.commands.update(
                {k: v.replace("\\n", "\n") for k, v in extra.items()})
        except Exception as e:
            logger.warning(f"voice_commands: could not load {path}: {e}")

    def _compile(self):
        phrases = sorted(self.commands.keys(), key=len, reverse=True)
        escaped = [re.escape(p) for p in phrases]
        pat = r'(?<![A-Za-z])(' + '|'.join(escaped) + r')(?![A-Za-z])'
        return re.compile(pat, re.IGNORECASE)

    def process(self, text: str) -> tuple:
        applied: list = []
        def replace(m):
            key = m.group(1).lower()
            applied.append(key)
            return self.commands.get(key, m.group(1))
        return self._pattern.sub(replace, text), applied

    def save(self, path: Path):
        data = {k: v.replace("\n", "\\n") for k, v in self.commands.items()}
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)


# ── Voice command editor ──────────────────────────────────────────────────────

class VoiceCommandEditor(tk.Toplevel):
    def __init__(self, parent, processor: VoiceCommandProcessor, on_save):
        super().__init__(parent)
        self.title("Voice Command Editor")
        self.configure(bg=BG)
        self.geometry("720x520")
        self.minsize(560, 380)
        self.resizable(True, True)
        self.transient(parent)
        self.grab_set()
        self._proc    = processor
        self._on_save = on_save
        self._data    = {k: v.replace("\n", "\\n")
                         for k, v in processor.commands.items()}
        self._build()
        self._populate()

    def _build(self):
        root = ttk.Frame(self, padding=16)
        root.pack(fill="both", expand=True)
        root.columnconfigure(0, weight=1)
        root.rowconfigure(2, weight=1)

        tk.Label(root, text="Voice Commands",
                 font=("Segoe UI", 13, "bold"), bg=BG, fg=ACCENT)\
            .grid(row=0, column=0, sticky="w")
        tk.Label(root,
                 text="Say the command word while dictating — it will be replaced.  "
                      "Use \\n for newline.",
                 font=FONT_SMALL, bg=BG, fg=TEXT_DIM)\
            .grid(row=1, column=0, sticky="w", pady=(2, 10))

        tf = tk.Frame(root, bg=BORDER, bd=1)
        tf.grid(row=2, column=0, sticky="nsew")
        tf.columnconfigure(0, weight=1)
        tf.rowconfigure(0, weight=1)

        self._tree = ttk.Treeview(tf, columns=("spoken", "output"),
                                   show="headings", selectmode="browse")
        self._tree.heading("spoken", text="You say (during dictation)")
        self._tree.heading("output", text="Output / formatting")
        self._tree.column("spoken", width=280, anchor="w", minwidth=160)
        self._tree.column("output", width=340, anchor="w", minwidth=160)
        vsb = ttk.Scrollbar(tf, orient="vertical", command=self._tree.yview)
        self._tree.configure(yscrollcommand=vsb.set)
        self._tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")

        s = ttk.Style(self)
        s.configure("Cmd.Treeview",
                     background=BG2, foreground=TEXT,
                     fieldbackground=BG2, rowheight=28, font=FONT_UI, borderwidth=0)
        s.configure("Cmd.Treeview.Heading",
                     background=ACCENT_L, foreground=ACCENT,
                     font=("Segoe UI", 10, "bold"), relief="flat", padding=6)
        s.map("Cmd.Treeview",
              background=[("selected", ACCENT)],
              foreground=[("selected", "white")])
        self._tree.configure(style="Cmd.Treeview")
        self._tree.bind("<Double-1>", self._on_double_click)
        self._tree.bind("<Delete>",   lambda _: self._delete_row())

        add = tk.Frame(root, bg=BG)
        add.grid(row=3, column=0, sticky="ew", pady=(12, 0))
        add.columnconfigure(1, weight=1)
        add.columnconfigure(4, weight=1)

        def lbl(t):
            return tk.Label(add, text=t, font=FONT_SMALL, bg=BG, fg=TEXT_MED)

        lbl("You say:").grid(row=0, column=0, padx=(0, 6))
        self._e_spoken = tk.Entry(
            add, bg=BG2, fg=TEXT, insertbackground=TEXT,
            relief="solid", bd=1, font=FONT_UI, highlightthickness=0)
        self._e_spoken.grid(row=0, column=1, sticky="ew", ipady=5)
        tk.Label(add, text="→", font=("Segoe UI", 12, "bold"),
                 bg=BG, fg=ACCENT).grid(row=0, column=2, padx=12)
        lbl("Output:").grid(row=0, column=3, padx=(0, 6))
        self._e_output = tk.Entry(
            add, bg=BG2, fg=TEXT, insertbackground=TEXT,
            relief="solid", bd=1, font=FONT_UI, highlightthickness=0)
        self._e_output.grid(row=0, column=4, sticky="ew", ipady=5)
        mk_btn(add, "Add / Update", self._add_or_update,
               style="open", size="small").grid(row=0, column=5, padx=(12, 0))
        self._e_spoken.bind("<Return>", lambda _: self._e_output.focus())
        self._e_output.bind("<Return>", lambda _: self._add_or_update())

        bot = tk.Frame(root, bg=BG)
        bot.grid(row=4, column=0, sticky="ew", pady=(12, 0))
        bot.columnconfigure(1, weight=1)
        mk_btn(bot, "🗑  Delete Selected", self._delete_row,
               style="danger", size="small").grid(row=0, column=0)
        self._count_lbl = tk.Label(bot, text="", font=FONT_SMALL,
                                    bg=BG, fg=TEXT_DIM)
        self._count_lbl.grid(row=0, column=1, padx=14, sticky="w")
        mk_btn(bot, "✓  Save & Close", self._save_close,
               style="save", size="small").grid(row=0, column=2)
        mk_btn(bot, "Cancel", self.destroy,
               style="subtle", size="small").grid(row=0, column=3, padx=(8, 0))

    def _populate(self):
        for row in self._tree.get_children():
            self._tree.delete(row)
        for spoken, output in sorted(self._data.items()):
            self._tree.insert("", "end", values=(spoken, output))
        self._count_lbl.config(text=f"{len(self._data)} commands")

    def _add_or_update(self):
        spoken = self._e_spoken.get().strip().lower()
        output = self._e_output.get().strip()
        if not spoken or not output:
            messagebox.showwarning("Empty field",
                                   "Both fields are required.", parent=self)
            return
        self._data[spoken] = output
        self._populate()
        self._e_spoken.delete(0, "end")
        self._e_output.delete(0, "end")
        self._e_spoken.focus()

    def _delete_row(self):
        sel = self._tree.selection()
        if not sel:
            return
        spoken = self._tree.item(sel[0], "values")[0]
        if messagebox.askyesno("Delete", f"Remove command '{spoken}'?", parent=self):
            self._data.pop(spoken, None)
            self._populate()

    def _on_double_click(self, _=None):
        sel = self._tree.selection()
        if not sel:
            return
        spoken, output = self._tree.item(sel[0], "values")
        self._e_spoken.delete(0, "end"); self._e_spoken.insert(0, spoken)
        self._e_output.delete(0, "end"); self._e_output.insert(0, output)
        self._e_output.focus()

    def _save_close(self):
        try:
            self._proc.commands = {
                k: v.replace("\\n", "\n") for k, v in self._data.items()}
            self._proc._pattern = self._proc._compile()
            if self._proc.config_path:
                self._proc.save(self._proc.config_path)
            self._on_save()
            self.destroy()
        except Exception as e:
            messagebox.showerror("Save failed", str(e), parent=self)


# ── Rewrite preview dialog ────────────────────────────────────────────────────

class RewritePreviewDialog(tk.Toplevel):
    def __init__(self, parent, original: str, rewritten: str,
                 on_accept, on_reject):
        super().__init__(parent)
        self.title("AI Rewrite Preview")
        self.configure(bg=BG)
        self.geometry("940x540")
        self.minsize(700, 380)
        self.resizable(True, True)
        self.transient(parent)
        self.grab_set()
        self._on_accept = on_accept
        self._on_reject = on_reject
        self._build(original, rewritten)
        self.focus_set()
        self.bind("<Escape>", lambda _: self._reject())

    def _build(self, original: str, rewritten: str):
        root = ttk.Frame(self, padding=16)
        root.pack(fill="both", expand=True)
        root.columnconfigure(0, weight=1)
        root.columnconfigure(1, weight=1)
        root.rowconfigure(1, weight=1)

        hdr = tk.Frame(root, bg=BG)
        hdr.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        tk.Label(hdr, text="AI Rewrite Preview",
                 font=("Segoe UI", 13, "bold"), bg=BG, fg=ACCENT).pack(side="left")
        tk.Label(hdr,
                 text="  You may edit the rewritten text before accepting.",
                 font=FONT_SMALL, bg=BG, fg=TEXT_DIM).pack(side="left", padx=(8, 0))

        lf = tk.Frame(root, bg=BG)
        lf.grid(row=1, column=0, sticky="nsew", padx=(0, 6))
        lf.columnconfigure(0, weight=1)
        lf.rowconfigure(1, weight=1)
        tk.Label(lf, text="Original selected text",
                 font=("Segoe UI", 10, "bold"), bg=BG, fg=TEXT_MED)\
            .grid(row=0, column=0, sticky="w", pady=(0, 4))
        orig_box = scrolledtext.ScrolledText(
            lf, wrap="word", font=FONT_MONO,
            bg="#FFF9F0", fg=TEXT, relief="solid", bd=1,
            highlightthickness=0, padx=10, pady=8, state="normal")
        orig_box.grid(row=1, column=0, sticky="nsew")
        orig_box.insert("1.0", original)
        orig_box.config(state="disabled")

        rf = tk.Frame(root, bg=BG)
        rf.grid(row=1, column=1, sticky="nsew", padx=(6, 0))
        rf.columnconfigure(0, weight=1)
        rf.rowconfigure(1, weight=1)
        tk.Label(rf, text="Rewritten text",
                 font=("Segoe UI", 10, "bold"), bg=BG, fg=ACCENT)\
            .grid(row=0, column=0, sticky="w", pady=(0, 4))
        self._rewr_edit = scrolledtext.ScrolledText(
            rf, wrap="word", font=FONT_MONO,
            bg="#F0FFF4", fg=TEXT, insertbackground=ACCENT,
            relief="solid", bd=1, highlightthickness=0,
            padx=10, pady=8, undo=True, state="normal")
        self._rewr_edit.grid(row=1, column=0, sticky="nsew")
        self._rewr_edit.insert("1.0", rewritten)
        self._rewr_edit.focus_set()

        tk.Frame(root, bg=BORDER, height=1)\
            .grid(row=2, column=0, columnspan=2, sticky="ew", pady=(12, 8))

        btns = tk.Frame(root, bg=BG)
        btns.grid(row=3, column=0, columnspan=2, sticky="ew")
        mk_btn(btns, "✓  Accept Rewrite", self._accept,
               style="save").pack(side="left")
        mk_btn(btns, "✗  Reject — Keep Original", self._reject,
               style="danger").pack(side="left", padx=(8, 0))
        tk.Label(btns, text="Esc = reject",
                 font=FONT_SMALL, bg=BG, fg=TEXT_DIM).pack(side="right", padx=(0, 4))

    def _accept(self):
        self._on_accept(self._rewr_edit.get("1.0", "end-1c"))
        self.destroy()

    def _reject(self):
        self._on_reject()
        self.destroy()


# ── Main application window ───────────────────────────────────────────────────

class App(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title(f"PathDictate  v{APP_VERSION}  —  Untitled")
        self.configure(bg=BG)
        self.minsize(720, 600)
        self.geometry("960x740")

        _icon = Path(__file__).parent / "app_icon.ico"
        if _icon.exists():
            try:
                self.iconbitmap(str(_icon))
            except Exception:
                pass

        self.cfg           = PathologyDictationConfig()
        self._ui_q         : queue.Queue = queue.Queue()
        self.is_recording  = False
        self.is_loading    = False
        self.transcriber   : Optional[PathologyTranscriber] = None
        self._recorder     : Optional[AudioRecorder]        = None
        self._corrector    : Optional[TerminologyCorrector] = None
        self._stream_lock  = threading.Lock()
        self._live_text    = ""
        self._cursor_on    = False
        self._rec_t0       = 0.0
        self._font_size    = FONT_SIZE_DEFAULT

        # ── Document state ────────────────────────────────────────────────────
        self._doc_path     : Optional[str] = None
        self._doc_dirty    : bool = False
        self._loading_file : bool = False   # suppresses dirty flag during load
        self._autosave_id                 = None

        # Privacy warning (read once from persisted settings)
        self._privacy_save_warned: bool = getattr(
            self.cfg, "_privacy_save_warned", False)

        # Voice commands
        _vc_path = Path(__file__).parent / "config" / "voice_commands.json"
        self._voice_cmd = VoiceCommandProcessor(config_path=_vc_path)

        # GGUF rewriter (legacy, lazy)
        self._rewriter: Optional[LocalRewriter] = None

        # Ollama rewrite service
        self._rewrite_svc: Optional[RewriteService] = None
        self._init_ollama()

        self._style()
        self._build()
        self._populate_mics()
        self._poll()
        self._autosave_tick()          # start autosave loop

        self._set_status("Loading model…", AMBER)
        threading.Thread(target=self._load_model, daemon=True).start()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ── ttk style ─────────────────────────────────────────────────────────────

    def _style(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        s.configure(".",            background=BG, foreground=TEXT, font=FONT_UI)
        s.configure("TFrame",       background=BG)
        s.configure("TLabel",       background=BG, foreground=TEXT)
        s.configure("Dim.TLabel",   background=BG, foreground=TEXT_DIM,
                     font=FONT_SMALL)
        s.configure("TLabelframe",  background=BG2, relief="groove",
                     bordercolor=BORDER)
        s.configure("TLabelframe.Label",
                     background=BG2, foreground=ACCENT,
                     font=("Segoe UI", 10, "bold"))
        s.configure("TNotebook",    background=BG, borderwidth=0)
        s.configure("TNotebook.Tab",
                     background=BG3, foreground=TEXT_DIM,
                     padding=(16, 7), font=("Segoe UI", 10))
        s.map("TNotebook.Tab",
              background=[("selected", BG2)],
              foreground=[("selected", ACCENT)],
              font=[("selected", ("Segoe UI", 10, "bold"))])
        s.configure("TCombobox",
                     fieldbackground=BG2, background=BG2, foreground=TEXT,
                     selectbackground=ACCENT_L, selectforeground=TEXT,
                     arrowcolor=TEXT_MED, bordercolor=BORDER,
                     lightcolor=BORDER, darkcolor=BORDER)
        s.map("TCombobox",
              fieldbackground=[("readonly", BG2)],
              bordercolor=[("focus", ACCENT)])
        s.configure("TScrollbar",
                     background=BG3, troughcolor=BG3,
                     arrowcolor=TEXT_DIM, borderwidth=0, relief="flat")
        s.map("TScrollbar", background=[("active", BORDER)])

    # ── Menubar ───────────────────────────────────────────────────────────────

    def _build_menubar(self):
        def _m(parent):
            return tk.Menu(parent, tearoff=0, bg=BG2, fg=TEXT,
                           activebackground=ACCENT_L, activeforeground=TEXT,
                           relief="solid", bd=1, font=FONT_SMALL)

        mb = tk.Menu(self, bg=BG2, fg=TEXT,
                     activebackground=ACCENT, activeforeground="white",
                     relief="flat", bd=0, font=FONT_UI)

        # File
        fm = _m(mb)
        fm.add_command(label="  📄  New Draft",        command=self._new_draft,
                       accelerator="Ctrl+N")
        fm.add_command(label="  📂  Open Text File…",  command=self._open_file,
                       accelerator="Ctrl+O")
        fm.add_separator()
        fm.add_command(label="  💾  Save",             command=self._save_file,
                       accelerator="Ctrl+S")
        fm.add_command(label="  💾  Save As…",         command=self._save_as,
                       accelerator="Ctrl+Shift+S")
        fm.add_separator()
        fm.add_command(label="  ✕  Close Draft",       command=self._close_draft)
        fm.add_separator()
        fm.add_command(label="  📄  Export as .docx",  command=self._save_docx)
        fm.add_separator()
        fm.add_command(label="  ⏻   Exit",              command=self._on_close)
        mb.add_cascade(label="File", menu=fm)

        # Edit
        em = _m(mb)
        em.add_command(label="  Undo",               command=self._undo,
                       accelerator="Ctrl+Z")
        em.add_command(label="  Redo",               command=self._redo,
                       accelerator="Ctrl+Y")
        em.add_separator()
        em.add_command(label="  📋  Copy to Clipboard", command=self._copy_result)
        em.add_separator()
        em.add_command(label="  ✕  Clear All",       command=self._clear_all)
        mb.add_cascade(label="Edit", menu=em)

        # Rewrite
        rm = _m(mb)
        rm.add_command(label="  ✏  Rewrite Selected Text",
                       command=self._rewrite_selected,
                       accelerator="Ctrl+Shift+R",
                       state="disabled")
        self._rewrite_menu = rm
        mb.add_cascade(label="Rewrite", menu=rm)

        # Tools
        tm = _m(mb)
        tm.add_command(label="  📖  Terminology Editor…",
                       command=self._open_editor)
        tm.add_command(label="  🎙  Voice Commands Editor…",
                       command=self._open_voice_editor)
        mb.add_cascade(label="Tools", menu=tm)

        # Settings
        sm = _m(mb)
        wm = _m(sm)
        for mdl in ["tiny", "base", "small", "medium", "large-v2", "large-v3"]:
            wm.add_command(label=f"  {mdl}",
                           command=lambda m=mdl: self._change_model_to(m))
        sm.add_cascade(label="  Whisper Model (named)", menu=wm)
        sm.add_command(label="  📂  Whisper Model Folder…",
                       command=self._browse_model_folder)
        sm.add_separator()
        fsm = _m(sm)
        fsm.add_command(label="  Larger  A+",  command=self._font_up)
        fsm.add_command(label="  Smaller A−",  command=self._font_down)
        fsm.add_command(label="  Reset",       command=self._font_reset)
        sm.add_cascade(label="  Font Size", menu=fsm)
        sm.add_separator()
        self._autocopy_var = tk.BooleanVar(value=self.cfg.ui.auto_copy_to_clipboard)
        sm.add_checkbutton(label="  Auto-copy to Clipboard",
                           variable=self._autocopy_var,
                           command=self._toggle_autocopy)
        sm.add_separator()
        sm.add_command(label="  Configure Ollama / AI Model…",
                       command=self._configure_ollama)
        mb.add_cascade(label="Settings", menu=sm)

        # Help
        hm = _m(mb)
        hm.add_command(label="  ❓  How to Use",         command=self._show_help)
        hm.add_command(label="  ⌨   Keyboard Shortcuts", command=self._show_shortcuts)
        hm.add_separator()
        hm.add_command(label="  ℹ   About",              command=self._show_about)
        mb.add_cascade(label="Help", menu=hm)

        self.config(menu=mb)

    # ── Settings helpers ──────────────────────────────────────────────────────

    def _change_model_to(self, model_name: str):
        self._mdl_var.set(model_name)
        self.cfg.whisper_model_path_override = ""
        self._update_model_label("loading")
        self._on_model_change()

    def _font_reset(self):
        self._font_size = FONT_SIZE_DEFAULT
        self._apply_font()

    def _toggle_autocopy(self):
        self.cfg.ui.auto_copy_to_clipboard = self._autocopy_var.get()

    def _configure_ollama(self):
        d = tk.Toplevel(self)
        d.title("Configure Ollama / AI Model")
        d.configure(bg=BG)
        d.geometry("540x330")
        d.minsize(420, 280)
        d.resizable(True, True)
        d.transient(self)
        d.grab_set()

        pad = ttk.Frame(d, padding=20)
        pad.pack(fill="both", expand=True)
        pad.columnconfigure(1, weight=1)

        tk.Label(pad, text="Ollama Configuration",
                 font=("Segoe UI", 13, "bold"), bg=BG, fg=ACCENT)\
            .grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 14))

        ep_var  = tk.StringVar(value=self.cfg.llm.endpoint)
        mdl_var = tk.StringVar(value=self.cfg.llm.model)
        to_var  = tk.StringVar(value=str(self.cfg.llm.timeout_seconds))

        for r, lbl_text, var in [
            (1, "Endpoint URL:",      ep_var),
            (2, "Model name:",        mdl_var),
            (3, "Timeout (seconds):", to_var),
        ]:
            tk.Label(pad, text=lbl_text, font=FONT_UI, bg=BG, fg=TEXT_MED)\
                .grid(row=r, column=0, sticky="w", padx=(0, 14), pady=4)
            tk.Entry(pad, textvariable=var, bg=BG2, fg=TEXT,
                     insertbackground=TEXT, relief="solid", bd=1,
                     font=FONT_UI, highlightthickness=0)\
                .grid(row=r, column=1, sticky="ew", ipady=5)

        status_lbl = tk.Label(pad, text="", font=FONT_SMALL,
                               bg=BG, fg=TEXT_DIM, anchor="w")
        status_lbl.grid(row=4, column=0, columnspan=2, sticky="ew", pady=(10, 0))

        btns = tk.Frame(pad, bg=BG)
        btns.grid(row=5, column=0, columnspan=2, sticky="ew", pady=(14, 0))

        def _test():
            status_lbl.config(text="Testing connection…", fg=AMBER)
            d.update_idletasks()
            try:
                tmp = OllamaClient(endpoint=ep_var.get(),
                                   model=mdl_var.get(), timeout=8)
                ok, msg = tmp.check_status()
                status_lbl.config(
                    text=("✓  " if ok else "✗  ") + msg[:90],
                    fg=GREEN_OK if ok else RED_REC)
            except Exception as exc:
                status_lbl.config(text=f"Error: {exc}", fg=RED_REC)

        def _save():
            try:
                self.cfg.llm.endpoint        = ep_var.get().strip()
                self.cfg.llm.model           = mdl_var.get().strip()
                self.cfg.llm.timeout_seconds = int(to_var.get().strip())
                if self._rewrite_svc:
                    c = self._rewrite_svc.client
                    c.base_url  = self.cfg.llm.endpoint
                    c._gen_url  = f"{self.cfg.llm.endpoint}/api/generate"
                    c._tags_url = f"{self.cfg.llm.endpoint}/api/tags"
                    c.model     = self.cfg.llm.model
                    c.timeout   = self.cfg.llm.timeout_seconds
                d.destroy()
            except ValueError:
                status_lbl.config(
                    text="Timeout must be a whole number of seconds", fg=RED_REC)

        mk_btn(btns, "Test Connection", _test,
               style="default", size="small").pack(side="left", padx=(0, 6))
        mk_btn(btns, "✓  Save", _save,
               style="save", size="small").pack(side="left", padx=(0, 6))
        mk_btn(btns, "Cancel", d.destroy,
               style="subtle", size="small").pack(side="left")

    # ── Help dialogs ──────────────────────────────────────────────────────────

    def _show_about(self):
        d = tk.Toplevel(self)
        d.title("About — PathDictate")
        d.configure(bg=BG)
        d.geometry("480x400")
        d.resizable(False, False)
        d.transient(self)
        d.grab_set()
        pad = ttk.Frame(d, padding=28)
        pad.pack(fill="both", expand=True)
        tk.Label(pad, text="🔬", font=("Segoe UI Emoji", 36), bg=BG).pack()
        tk.Label(pad, text="PathDictate",
                 font=("Segoe UI", 15, "bold"), bg=BG, fg=ACCENT).pack(pady=(6, 2))
        tk.Label(pad,
                 text=f"Version {APP_VERSION}   •   Local · Offline · Privacy-first",
                 font=FONT_SMALL, bg=BG, fg=TEXT_DIM).pack()
        tk.Frame(pad, bg=BORDER, height=1).pack(fill="x", pady=16)
        tk.Label(pad,
                 text="Transcribes pathology dictations locally using OpenAI Whisper\n"
                      "(faster-whisper)  —  no cloud, no internet required.\n\n"
                      "AI rewrite uses Ollama + Qwen2.5 on localhost.\n"
                      "No patient data ever leaves this computer.",
                 font=FONT_UI, bg=BG, fg=TEXT_MED, justify="center").pack()
        tk.Frame(pad, bg=BORDER, height=1).pack(fill="x", pady=16)
        mk_btn(pad, "Close", d.destroy, style="open", size="small").pack(pady=(4, 0))

    def _show_help(self):
        d = tk.Toplevel(self)
        d.title("How to Use — PathDictate")
        d.configure(bg=BG)
        d.geometry("680x640")
        d.minsize(540, 460)
        d.resizable(True, True)
        d.transient(self)
        d.grab_set()
        pad = ttk.Frame(d, padding=20)
        pad.pack(fill="both", expand=True)
        pad.rowconfigure(1, weight=1)
        pad.columnconfigure(0, weight=1)
        tk.Label(pad, text="How to Use",
                 font=("Segoe UI", 13, "bold"), bg=BG, fg=ACCENT)\
            .grid(row=0, column=0, sticky="w", pady=(0, 10))
        txt = scrolledtext.ScrolledText(
            pad, wrap="word", font=("Segoe UI", 11),
            bg=BG2, fg=TEXT, relief="solid", bd=1,
            padx=14, pady=10, highlightthickness=0)
        txt.grid(row=1, column=0, sticky="nsew")
        txt.insert("1.0", """\
BASIC DICTATION
━━━━━━━━━━━━━━━
  1. Press F9 (or ⏺ Start Recording)
  2. Speak clearly
  3. Press F9 again to stop
  4. Text is inserted at cursor position in the Corrected panel
  5. Corrected text is auto-copied to clipboard

OPEN / SAVE DRAFTS
━━━━━━━━━━━━━━━━━━
  Ctrl+N          New draft  (warns if unsaved changes)
  Ctrl+O          Open a .txt draft file
  Ctrl+S          Save  (saves to current file or prompts for new name)
  Ctrl+Shift+S    Save As  (always prompts for name)
  File > Close    Close current draft
  File > Export   Export as .docx (requires python-docx)

  Default filename: PathDictate_Draft_YYYYMMDD_HHMMSS.txt
  Unsaved changes shown as  *  in the window title.

CONTINUE DICTATION AT CURSOR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Click anywhere in the Corrected panel to place the cursor.
  The next dictation inserts text at that position.
  At end-of-document: a  ― ― ―  separator is added first.
  Mid-document: text is inserted directly at cursor (no separator).

AUTOSAVE
━━━━━━━━
  Runs every 60 seconds when there are unsaved changes.
  Backup saved to:  ./autosave/PathDictate_autosave.txt
  Does NOT overwrite your named file — it is a backup only.

AI REWRITE  (requires Ollama)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  1. Select text in the Corrected panel
  2. Click  ✏ Rewrite Selected Text  or press Ctrl+Shift+R
  3. Review side-by-side preview — edit the rewritten text if needed
  4. Accept or Reject
  The AI is strictly instructed never to add new findings or diagnoses.
  Install model:  ollama pull qwen2.5:14b

WHISPER MODEL
━━━━━━━━━━━━━
  Settings > Whisper Model (named)    — choose tiny/base/small/medium/large
  Settings > Whisper Model Folder     — browse to a local faster-whisper folder
  Header shows current model path.    ● green = loaded,  ● red = not found.
  On portable machines: point to your local model folder using Browse.

UNDO / REDO
━━━━━━━━━━━
  Ctrl+Z  Undo    Ctrl+Y  Redo
  Works for manual edits, dictation insertions, and accepted rewrites.
""")
        txt.config(state="disabled")
        mk_btn(pad, "Close", d.destroy, style="open",
               size="small").grid(row=2, column=0, pady=(12, 0))

    def _show_shortcuts(self):
        d = tk.Toplevel(self)
        d.title("Keyboard Shortcuts")
        d.configure(bg=BG)
        d.geometry("440x480")
        d.resizable(False, False)
        d.transient(self)
        d.grab_set()
        pad = ttk.Frame(d, padding=24)
        pad.pack(fill="both", expand=True)
        tk.Label(pad, text="Keyboard Shortcuts",
                 font=("Segoe UI", 13, "bold"), bg=BG, fg=ACCENT).pack(anchor="w")
        tk.Frame(pad, bg=BORDER, height=1).pack(fill="x", pady=(8, 16))

        shortcuts = [
            ("F9",            "Start / Stop recording"),
            ("Esc",           "Cancel recording  (no transcription)"),
            ("", ""),
            ("Ctrl+N",        "New draft"),
            ("Ctrl+O",        "Open text file"),
            ("Ctrl+S",        "Save"),
            ("Ctrl+Shift+S",  "Save As"),
            ("", ""),
            ("Ctrl+Z",        "Undo last edit"),
            ("Ctrl+Y",        "Redo"),
            ("", ""),
            ("Ctrl+Shift+R",  "Rewrite selected text with AI"),
        ]
        for key, desc in shortcuts:
            if not key and not desc:
                tk.Frame(pad, bg=BG, height=6).pack(fill="x")
                continue
            row = tk.Frame(pad, bg=BG)
            row.pack(fill="x", pady=2)
            tk.Label(row, text=key,
                     font=("Consolas", 10, "bold"),
                     bg=ACCENT_L, fg=ACCENT, padx=10, pady=3,
                     width=18, anchor="center").pack(side="left")
            tk.Label(row, text=desc, font=FONT_UI,
                     bg=BG, fg=TEXT).pack(side="left", padx=(12, 0))

        tk.Frame(pad, bg=BORDER, height=1).pack(fill="x", pady=(16, 12))
        mk_btn(pad, "Close", d.destroy, style="open", size="small").pack()

    # ── Build UI ──────────────────────────────────────────────────────────────

    def _build(self):
        self._build_menubar()

        outer = ttk.Frame(self, padding=14)
        outer.pack(fill="both", expand=True)
        outer.columnconfigure(0, weight=1)
        outer.rowconfigure(2, weight=1)

        self._build_header(outer)
        self._build_recbar(outer)
        self._build_notebook(outer)
        self._build_bottom(outer)

        # Keyboard shortcuts
        self.bind("<F9>",              lambda _: self._toggle_record())
        self.bind("<Escape>",          lambda _:
                  self._stop_only() if self.is_recording else None)
        self.bind("<Control-Shift-R>", lambda _: self._rewrite_selected())
        self.bind("<Control-Shift-r>", lambda _: self._rewrite_selected())
        self.bind("<Control-n>",       lambda _: self._new_draft())
        self.bind("<Control-N>",       lambda _: self._new_draft())
        self.bind("<Control-o>",       lambda _: self._open_file())
        self.bind("<Control-O>",       lambda _: self._open_file())
        self.bind("<Control-s>",       lambda _: self._save_file())
        self.bind("<Control-S>",       lambda _: self._save_file())
        self.bind("<Control-Shift-s>", lambda _: self._save_as())
        self.bind("<Control-Shift-S>", lambda _: self._save_as())

    # ── Header ────────────────────────────────────────────────────────────────

    def _build_header(self, parent):
        hdr = ttk.Frame(parent)
        hdr.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        hdr.columnconfigure(1, weight=1)

        tk.Label(hdr, text="🔬  PathDictate",
                 font=FONT_HEAD, bg=BG, fg=ACCENT)\
            .grid(row=0, column=0, sticky="w")
        tk.Label(hdr,
                 text="Local · Offline · Privacy-first · v" + APP_VERSION,
                 font=FONT_SMALL, bg=BG, fg=TEXT_DIM)\
            .grid(row=1, column=0, sticky="w", pady=(2, 0))

        # Right-side controls
        cfg_frame = ttk.Frame(hdr)
        cfg_frame.grid(row=0, column=2, rowspan=3, sticky="e")

        # Labels
        tk.Label(cfg_frame, text="Model",
                 font=FONT_SMALL, bg=BG, fg=TEXT_MED)\
            .grid(row=0, column=0, padx=(0, 4), sticky="w")
        tk.Label(cfg_frame, text="Microphone",
                 font=FONT_SMALL, bg=BG, fg=TEXT_MED)\
            .grid(row=0, column=1, padx=(0, 4), sticky="w")

        # Comboboxes
        self._mdl_var = tk.StringVar(value=self.cfg.transcription.model_size)
        self._mdl_cb  = ttk.Combobox(
            cfg_frame, textvariable=self._mdl_var, width=12, state="readonly",
            values=["tiny", "base", "small", "medium", "large-v2", "large-v3"])
        self._mdl_cb.grid(row=1, column=0, padx=(0, 8))
        self._mdl_cb.bind("<<ComboboxSelected>>", self._on_model_change)

        self._mic_var = tk.StringVar()
        self._mic_cb  = ttk.Combobox(cfg_frame, textvariable=self._mic_var,
                                      width=26, state="readonly")
        self._mic_cb.grid(row=1, column=1)
        self._mic_cb.bind("<<ComboboxSelected>>", self._on_mic_change)

        # Model path status row
        mp_row = tk.Frame(cfg_frame, bg=BG)
        mp_row.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(6, 0))

        self._mdl_status_dot = tk.Label(
            mp_row, text="●", font=FONT_SMALL, bg=BG, fg=AMBER)
        self._mdl_status_dot.pack(side="left")

        self._mdl_path_lbl = tk.Label(
            mp_row, text="  Loading…",
            font=("Segoe UI", 9), bg=BG, fg=TEXT_DIM,
            anchor="w", width=32)
        self._mdl_path_lbl.pack(side="left", padx=(2, 8))

        mk_btn(mp_row, "Browse Model…", self._browse_model_folder,
               style="default", size="small").pack(side="left", padx=(0, 4))
        mk_btn(mp_row, "Reload", self._on_model_change,
               style="default", size="small").pack(side="left")

    # ── Record bar ────────────────────────────────────────────────────────────

    def _build_recbar(self, parent):
        card = tk.Frame(parent, bg=BG2,
                         highlightthickness=1, highlightbackground=BORDER)
        card.grid(row=1, column=0, sticky="ew", pady=(0, 10))
        card.columnconfigure(1, weight=1)

        self._rec_btn = mk_btn(
            card, "⏺   Start Recording   (F9)",
            self._toggle_record,
            style="record", size="large")
        self._rec_btn.grid(row=0, column=0, padx=16, pady=12)

        right = tk.Frame(card, bg=BG2)
        right.grid(row=0, column=1, sticky="nsew", padx=(0, 16), pady=10)
        right.columnconfigure(0, weight=1)

        top_row = tk.Frame(right, bg=BG2)
        top_row.grid(row=0, column=0, sticky="ew")
        top_row.columnconfigure(0, weight=1)

        self._stat_lbl = tk.Label(
            top_row, text="Initializing…",
            font=("Segoe UI", 11, "bold"),
            bg=BG2, fg=AMBER, anchor="w")
        self._stat_lbl.grid(row=0, column=0, sticky="w")

        self._time_lbl = tk.Label(
            top_row, text="",
            font=("Segoe UI Semibold", 11),
            bg=BG2, fg=TEXT_DIM, anchor="e")
        self._time_lbl.grid(row=0, column=1, sticky="e")

        self._vu = VUMeter(right)
        self._vu.grid(row=1, column=0, sticky="ew", pady=(6, 0))

        tk.Label(right,
                 text="Press F9 to start · F9 again to stop · Esc to cancel",
                 font=FONT_SMALL, bg=BG2, fg=TEXT_DIM, anchor="w")\
            .grid(row=2, column=0, sticky="w", pady=(4, 0))

    # ── Notebook ──────────────────────────────────────────────────────────────

    def _build_notebook(self, parent):
        wrap = ttk.LabelFrame(parent, text="  Transcription  ", padding=(12, 8))
        wrap.grid(row=2, column=0, sticky="nsew", pady=(0, 8))
        wrap.columnconfigure(0, weight=1)
        wrap.rowconfigure(0, weight=1)

        self._nb = ttk.Notebook(wrap)
        self._nb.grid(row=0, column=0, sticky="nsew")

        # Live tab — readonly streaming preview
        self._live_box = self._textbox(self._nb)
        self._nb.add(self._live_box, text="  🎙  Live  ")

        self._corr_box = self._editable_textbox(self._nb)
        self._nb.add(self._corr_box, text="  ✏  Corrected  ")

        self._nb.select(0)

        # Dirty-state tracking and redo shortcut
        self._corr_box.bind("<<Modified>>", self._on_text_modified)
        self._corr_box.bind("<Control-y>",  lambda e: self._redo() or "break")
        self._corr_box.bind("<Control-Y>",  lambda e: self._redo() or "break")

        self._build_edit_toolbar(wrap)

        self._changes_lbl = tk.Label(
            wrap, text="", font=("Segoe UI", 10), bg=BG2,
            fg=TEXT_MED, anchor="w", justify="left", wraplength=880)
        self._changes_lbl.grid(row=2, column=0, sticky="w", pady=(4, 0))

    def _textbox(self, parent) -> scrolledtext.ScrolledText:
        return scrolledtext.ScrolledText(
            parent, wrap="word",
            font=("Segoe UI", self._font_size),
            height=1,
            bg=BG2, fg=TEXT, insertbackground=ACCENT,
            relief="flat", bd=0, padx=14, pady=12,
            selectbackground=ACCENT_L, selectforeground=TEXT,
            state="disabled")

    def _editable_textbox(self, parent) -> scrolledtext.ScrolledText:
        return scrolledtext.ScrolledText(
            parent, wrap="word",
            font=("Segoe UI", self._font_size),
            height=1,
            bg="#F7FFFE", fg=TEXT, insertbackground=ACCENT,
            relief="flat", bd=0, padx=14, pady=12,
            selectbackground=ACCENT_L, selectforeground=TEXT,
            undo=True, maxundo=200,
            state="normal")

    def _build_edit_toolbar(self, parent):
        bar = tk.Frame(parent, bg=BG2)
        bar.grid(row=1, column=0, sticky="ew", pady=(5, 0))

        mk_btn(bar, "↩  Undo  Ctrl+Z", self._undo,
               style="subtle", size="small").pack(side="left", padx=(0, 2))
        mk_btn(bar, "↪  Redo  Ctrl+Y", self._redo,
               style="subtle", size="small").pack(side="left", padx=(0, 2))

        tk.Frame(bar, bg=BORDER, width=1)\
            .pack(side="left", fill="y", padx=(8, 8), pady=3)

        self._rewrite_sel_btn = mk_btn(
            bar, "✏  Rewrite Selected Text   Ctrl+Shift+R",
            self._rewrite_selected,
            style="rewrite", size="small",
            state="disabled")
        self._rewrite_sel_btn.pack(side="left")

        tk.Label(bar,
                 text="✎ Corrected tab is editable — "
                      "click to place cursor, then dictate to insert there",
                 font=("Segoe UI", 9), bg=BG2, fg=TEXT_DIM)\
            .pack(side="right", padx=(0, 6))

    def _editable_textbox(self, parent) -> scrolledtext.ScrolledText:
        """Textbox that stays editable — undo/redo enabled, user can type freely."""
        return scrolledtext.ScrolledText(
            parent, wrap="word",
            font=("Segoe UI", self._font_size),
            height=1,
            bg="#F7FFFE", fg=TEXT, insertbackground=ACCENT,
            relief="flat", bd=0, padx=14, pady=12,
            selectbackground=ACCENT_L, selectforeground=TEXT,
            undo=True, maxundo=200,
            state="normal"
        )

    def _build_edit_toolbar(self, parent):
        """
        Compact toolbar below the notebook tabs.
        Contains: Undo · Redo | Rewrite Selected Text
        This toolbar only operates on the Corrected (editable) panel.
        """
        bar = tk.Frame(parent, bg=BG2)
        bar.grid(row=1, column=0, sticky="ew", pady=(5, 0))

        def tbtn(text, cmd, fg=TEXT_MED, **kw):
            b = tk.Button(bar, text=text, font=("Segoe UI", 9),
                          bg=BG2, fg=fg,
                          activebackground=ACCENT_L, activeforeground=ACCENT,
                          relief="flat", bd=0, padx=10, pady=4,
                          cursor="hand2", command=cmd, **kw)
            b.pack(side="left", padx=(0, 2))
            return b

        tbtn("↩  Undo  Ctrl+Z", self._undo)
        tbtn("↪  Redo  Ctrl+Y", self._redo)

        tk.Frame(bar, bg=BORDER, width=1)\
            .pack(side="left", fill="y", padx=(8, 8), pady=3)

        self._rewrite_sel_btn = tbtn(
            "✏  Rewrite Selected Text   Ctrl+Shift+R",
            self._rewrite_selected, fg=ACCENT,
            state="disabled")   # enabled once Ollama startup check passes

        # Right-side hint
        tk.Label(bar,
                 text="✎ Corrected tab is editable — type freely or rewrite selection with AI",
                 font=("Segoe UI", 9), bg=BG2, fg=TEXT_DIM)\
            .pack(side="right", padx=(0, 6))

    # ── Bottom bar ────────────────────────────────────────────────────────────

    def _build_bottom(self, parent):
        bar = tk.Frame(parent, bg=BG2,
                        highlightthickness=1, highlightbackground=BORDER)
        bar.grid(row=3, column=0, sticky="ew")

        row0 = tk.Frame(bar, bg=BG2)
        row0.pack(fill="x", padx=8, pady=(7, 3))

        def sep_l():
            tk.Frame(row0, bg=BORDER, width=1).pack(
                side="left", fill="y", padx=(6, 6), pady=4)

        def sep_r():
            tk.Frame(row0, bg=BORDER, width=1).pack(
                side="right", fill="y", padx=(6, 6), pady=4)

        # Right side
        mk_btn(row0, "⏻  Exit", self._on_close,
               style="danger", size="small").pack(side="right")
        sep_r()
        mk_btn(row0, "A+", self._font_up,   size="small").pack(side="right")
        self._size_lbl = tk.Label(
            row0, text=str(self._font_size), width=3,
            font=("Segoe UI", 10, "bold"),
            bg=BG2, fg=TEXT_MED, anchor="center")
        self._size_lbl.pack(side="right")
        mk_btn(row0, "A−", self._font_down, size="small").pack(side="right")
        tk.Label(row0, text="Size:", font=FONT_SMALL,
                 bg=BG2, fg=TEXT_DIM).pack(side="right", padx=(0, 2))
        sep_r()

        # Left side
        mk_btn(row0, "📖  Terminology",  self._open_editor,
               style="open",    size="small").pack(side="left", padx=(0, 3))
        mk_btn(row0, "🎙  Voice Cmds",  self._open_voice_editor,
               style="default", size="small").pack(side="left", padx=(0, 3))
        sep_l()
        mk_btn(row0, "📂  Open",        self._open_file,
               style="open",    size="small").pack(side="left", padx=(0, 3))
        self._save_btn = mk_btn(row0, "💾  Save", self._save_file,
                                 style="save", size="small")
        self._save_btn.pack(side="left", padx=(0, 3))
        mk_btn(row0, "💾  Save As",     self._save_as,
               style="save",    size="small").pack(side="left", padx=(0, 3))
        sep_l()
        mk_btn(row0, "📋  Copy",        self._copy_result,
               style="default", size="small").pack(side="left", padx=(0, 3))
        mk_btn(row0, "✕  Clear",        self._clear_all,
               style="subtle",  size="small").pack(side="left")

        # Footer status label
        self._foot = tk.Label(bar, text="", font=FONT_SMALL,
                               bg=BG2, fg=TEXT_DIM, anchor="w")
        self._foot.pack(fill="x", padx=10, pady=(0, 5))

    # ── Model path helpers ────────────────────────────────────────────────────

    def _update_model_label(self, status: str = "pending"):
        """Update the status dot and path label in the header."""
        override = self.cfg.whisper_model_path_override
        shown    = (Path(override).name or override) if override \
                   else self._mdl_var.get()
        if len(shown) > 34:
            shown = "…" + shown[-32:]

        dot_color = {
            "loading": AMBER,   "ok":      GREEN_OK,
            "error":   RED_REC, "missing": RED_REC,
            "pending": TEXT_DIM,
        }.get(status, TEXT_DIM)

        self._mdl_status_dot.config(fg=dot_color)
        self._mdl_path_lbl.config(text=f"  {shown}")

    def _browse_model_folder(self):
        """Open a directory picker for the local faster-whisper model folder."""
        start = (self.cfg.whisper_model_path_override
                 or str(self.cfg.models_dir))
        folder = filedialog.askdirectory(
            title="Select faster-whisper model folder",
            initialdir=start,
            mustexist=True)
        if not folder:
            return

        p = Path(folder)
        if not self._validate_model_folder(p):
            if not messagebox.askyesno(
                "Folder may not be valid",
                f"This folder does not appear to contain a faster-whisper model.\n\n"
                f"Expected files such as:  model.bin,  config.json,  tokenizer.json\n\n"
                f"Folder: {folder}\n\n"
                f"Try loading it anyway?",
                parent=self):
                return

        self.cfg.whisper_model_path_override = str(p)
        self.cfg.save_user_settings()
        self._update_model_label("loading")
        self._on_model_change()

    def _validate_model_folder(self, path: Path) -> bool:
        """Light check that a folder looks like a faster-whisper/CTranslate2 model."""
        if not path.is_dir():
            return False
        files = {f.name.lower() for f in path.iterdir() if f.is_file()}
        has_config = "config.json" in files
        has_model  = any(
            n.endswith((".bin", ".ot", ".onnx")) or n.startswith("model")
            for n in files)
        return has_config or has_model

    # ── Mic list ──────────────────────────────────────────────────────────────

    def _populate_mics(self):
        self._mic_map = {}
        try:
            for i, dev in enumerate(sd.query_devices()):
                if dev["max_input_channels"] > 0:
                    name = dev["name"][:42]
                    self._mic_map[name] = i
            names = list(self._mic_map)
            self._mic_cb["values"] = names
            if names:
                default = sd.default.device
                def_idx = default[0] if isinstance(default, (list, tuple)) \
                          else default
                pick = next((n for n, i in self._mic_map.items()
                             if i == def_idx), names[0])
                self._mic_var.set(pick)
                self.cfg.audio.device_index = self._mic_map[pick]
        except Exception as e:
            logger.warning(f"mic list: {e}")

    def _on_mic_change(self, _=None):
        self.cfg.audio.device_index = self._mic_map.get(self._mic_var.get())
        self._recorder = None

    # ── Model loading ─────────────────────────────────────────────────────────

    def _load_model(self):
        self.is_loading = True
        self._btn_enabled(False)
        self._ui_q.put(("model_status", "loading", ""))
        try:
            # Resolve effective model path
            override = self.cfg.whisper_model_path_override
            if override:
                model_path = override
            else:
                model_path = self._mdl_var.get()

            # Resolve relative paths
            if model_path.startswith(("./", ".\\")):
                model_path = str(Path(__file__).parent / model_path[2:])

            self.cfg.transcription.model_size = model_path
            self.cfg.transcription.device     = "auto"   # CUDA auto-detect

            self.transcriber = PathologyTranscriber(
                self.cfg.transcription, self.cfg.models_dir)
            self._corrector  = TerminologyCorrector(
                self.cfg.dictionary, self.cfg.dictionary_path)

            disp = Path(model_path).name if Path(model_path).exists() \
                   else model_path
            self._ui_q.put(("status",       f"Ready  ·  model: {disp}", GREEN_OK))
            self._ui_q.put(("model_status", "ok",   disp))

        except FileNotFoundError:
            msg = (f"Model folder not found:\n{self.cfg.transcription.model_size}\n\n"
                   "Use Settings > Whisper Model Folder to select a local folder.")
            self._ui_q.put(("status",       "Model not found — select a local folder",
                             RED_REC))
            self._ui_q.put(("model_status", "missing", ""))
            self._ui_q.put(("msgbox",       "error", "Model not found", msg))
        except Exception as e:
            self._ui_q.put(("status",       f"Model error: {e}", RED_REC))
            self._ui_q.put(("model_status", "error", ""))
            self._ui_q.put(("msgbox",       "error", "Model failed",
                             f"Could not load model:\n{e}"))
        finally:
            self.is_loading = False
            self._ui_q.put(("btn", True))

    def _on_model_change(self, _=None):
        if self.is_recording:
            return
        self._set_status("Reloading model…", AMBER)
        self._btn_enabled(False)
        threading.Thread(target=self._load_model, daemon=True).start()

    # ── Terminology / Voice editors ───────────────────────────────────────────

    def _open_editor(self):
        def _reload():
            if self._corrector:
                self._corrector._load_dictionary()
                n = len(self._corrector.replacements)
                self._set_status(
                    f"Dictionary reloaded  ·  {n} entries", GREEN_OK)
        TerminologyEditor(self, self.cfg.dictionary_path, _reload)

    def _open_voice_editor(self):
        def _reload():
            n = len(self._voice_cmd.commands)
            self._set_status(
                f"Voice commands saved  ·  {n} commands active", GREEN_OK)
        VoiceCommandEditor(self, self._voice_cmd, _reload)

    # ── Undo / Redo ───────────────────────────────────────────────────────────

    def _undo(self):
        try:
            self._corr_box.edit_undo()
        except tk.TclError:
            pass

    def _redo(self):
        try:
            self._corr_box.edit_redo()
        except tk.TclError:
            pass

    # ── Document management ───────────────────────────────────────────────────

    def _mark_dirty(self):
        if not self._loading_file:
            self._doc_dirty = True
            self._update_title()

    def _update_title(self):
        name  = Path(self._doc_path).name if self._doc_path else "Untitled"
        dirty = " *" if self._doc_dirty else ""
        self.title(f"PathDictate  v{APP_VERSION}  —  {name}{dirty}")

    def _on_text_modified(self, _=None):
        """Fire whenever the Corrected widget changes (manual typing or insertion)."""
        self._corr_box.edit_modified(False)   # reset so event fires again next time
        self._mark_dirty()

    def _check_dirty_before_action(self, callback):
        """Prompt user to save if dirty, then run callback."""
        if not self._doc_dirty:
            callback()
            return
        ans = messagebox.askyesnocancel(
            "Unsaved changes",
            "You have unsaved changes.\nSave before continuing?",
            parent=self)
        if ans is None:
            return
        if ans:
            if not self._save_file():
                return
        callback()

    def _new_draft(self):
        def _do():
            self._loading_file = True
            self._corr_box.delete("1.0", "end")
            self._write_box(self._live_box, "")
            self._changes_lbl.config(text="")
            self._live_text = ""
            self._doc_path  = None
            self._doc_dirty = False
            self._loading_file = False
            self._corr_box.edit_modified(False)
            self._update_title()
            self._set_status("New draft started", GREEN_OK)
        self._check_dirty_before_action(_do)

    def _open_file(self):
        def _do():
            start = (self.cfg.documents.last_draft_folder
                     or str(self.cfg.drafts_dir))
            path = filedialog.askopenfilename(
                title="Open draft",
                initialdir=start,
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
            if not path:
                return
            try:
                with open(path, encoding="utf-8") as f:
                    content = f.read()
            except Exception as e:
                messagebox.showerror("Open failed", str(e), parent=self)
                return

            self._loading_file = True
            self._corr_box.delete("1.0", "end")
            self._corr_box.insert("1.0", content)
            self._corr_box.edit_modified(False)
            self._corr_box.mark_set(tk.INSERT, "end-1c")
            self._corr_box.see("end")
            self._loading_file = False

            self._doc_path  = path
            self._doc_dirty = False
            self.cfg.documents.last_draft_folder = str(Path(path).parent)
            self.cfg.save_user_settings()
            self._update_title()
            self._nb.select(1)
            self._set_status(f"Opened: {Path(path).name}", GREEN_OK)
        self._check_dirty_before_action(_do)

    def _save_file(self) -> bool:
        """Save to current path; falls back to Save As. Returns True on success."""
        if self._doc_path:
            return self._write_file(self._doc_path)
        return self._save_as()

    def _save_as(self) -> bool:
        """Prompt for filename and save. Returns True on success."""
        if not self._corr_box.get("1.0", "end-1c").strip():
            messagebox.showwarning("Nothing to save",
                                   "The editor is empty.", parent=self)
            return False
        stamp   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        default = f"{self.cfg.documents.default_filename_prefix}_{stamp}.txt"
        start   = (self.cfg.documents.last_draft_folder
                   or str(self.cfg.drafts_dir))
        path = filedialog.asksaveasfilename(
            title="Save draft as",
            initialdir=start,
            initialfile=default,
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
        if not path:
            return False
        if self._write_file(path):
            self._doc_path = path
            self.cfg.documents.last_draft_folder = str(Path(path).parent)
            self.cfg.save_user_settings()
            return True
        return False

    def _write_file(self, path: str) -> bool:
        """Write corrected-box content to path. Returns True on success."""
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self._corr_box.get("1.0", "end-1c"))
            self._doc_dirty = False
            self._update_title()
            self._set_status(f"Saved: {Path(path).name}", GREEN_OK)
            self._show_privacy_note()
            return True
        except Exception as e:
            messagebox.showerror("Save failed", str(e), parent=self)
            return False

    def _close_draft(self):
        def _do():
            self._loading_file = True
            self._corr_box.delete("1.0", "end")
            self._write_box(self._live_box, "")
            self._changes_lbl.config(text="")
            self._live_text = ""
            self._doc_path  = None
            self._doc_dirty = False
            self._loading_file = False
            self._corr_box.edit_modified(False)
            self._update_title()
            self._set_status("Draft closed", TEXT_DIM)
        self._check_dirty_before_action(_do)

    # ── Autosave ──────────────────────────────────────────────────────────────

    def _autosave_tick(self):
        interval_ms = self.cfg.documents.autosave_interval_seconds * 1000
        if (self.cfg.documents.autosave_enabled
                and self._doc_dirty
                and self._corr_box.get("1.0", "end-1c").strip()):
            self._do_autosave()
        self._autosave_id = self.after(interval_ms, self._autosave_tick)

    def _do_autosave(self):
        try:
            path = self.cfg.autosave_dir / "PathDictate_autosave.txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(self._corr_box.get("1.0", "end-1c"))
            logger.debug(f"Autosaved to {path}")
            self._show_privacy_note()
        except Exception as e:
            logger.warning(f"Autosave failed: {e}")

    # ── Privacy note (shown once) ─────────────────────────────────────────────

    def _show_privacy_note(self):
        if self._privacy_save_warned:
            return
        self._privacy_save_warned     = True
        self.cfg._privacy_save_warned = True
        self.cfg.save_user_settings()
        messagebox.showinfo(
            "Privacy reminder",
            "Saved drafts may contain patient-identifiable information.\n\n"
            "Save only to an approved, secure local folder.\n\n"
            "(This message is shown once only.)",
            parent=self)

    # ── Insert dictation at cursor ────────────────────────────────────────────

    def _insert_at_cursor_in_corr(self, text: str):
        """
        Insert new dictation text at the current cursor position.
        · Empty box          → insert at start.
        · Cursor at end      → add  ― ― ―  separator then text.
        · Cursor mid-document → insert directly at cursor (continuation mode).
        """
        box      = self._corr_box
        existing = box.get("1.0", "end-1c")
        cursor   = box.index(tk.INSERT)
        end_idx  = box.index("end-1c")

        if not existing.strip():
            box.insert("1.0", text)
        elif cursor >= end_idx:
            # at or past end — add separator
            sep = "\n\n― ― ―\n\n"
            sep_start = box.index("end-1c")
            box.insert("end", sep)
            box.tag_add("sep", sep_start, box.index("end-1c"))
            box.tag_config("sep", foreground=TEXT_DIM)
            box.insert("end", text)
        else:
            # mid-document — insert directly
            box.insert(cursor, text)

        box.edit_separator()   # single undo point
        box.see(tk.INSERT)
        self._mark_dirty()

    # ── Ollama / rewrite ──────────────────────────────────────────────────────

    def _init_ollama(self):
        cfg = self.cfg.llm
        if not cfg.enabled:
            return
        try:
            client = OllamaClient(
                endpoint=cfg.endpoint,
                model=cfg.model,
                timeout=cfg.timeout_seconds)
            self._rewrite_svc = RewriteService(
                client=client,
                temperature=cfg.temperature,
                max_tokens=cfg.max_tokens)
        except Exception as e:
            logger.warning(f"Could not create Ollama service: {e}")
            self._rewrite_svc = None
            return
        threading.Thread(
            target=self._check_ollama_startup,
            daemon=True, name="ollama-startup").start()

    def _check_ollama_startup(self):
        cfg    = self.cfg.llm
        client = self._rewrite_svc.client
        self._ui_q.put(("ollama_foot", "Ollama: checking…", TEXT_DIM))

        if client.is_ollama_running():
            self._finish_ollama_check(client, cfg)
            return

        if not cfg.auto_start_ollama:
            self._ui_q.put(("ollama_status", "unavailable",
                             "Ollama not running.  Start with:  ollama serve"))
            return

        self._ui_q.put(("ollama_foot", "Ollama: starting…", AMBER))
        launched = client.start_ollama(cfg.ollama_start_command)
        if not launched:
            self._ui_q.put(("ollama_status", "unavailable",
                             "Could not start Ollama.  Install from https://ollama.ai"))
            return

        self._ui_q.put(("ollama_foot",
                         f"Ollama: waiting up to {cfg.startup_wait_seconds} s…",
                         AMBER))
        ready = client.wait_until_ready(
            timeout=cfg.startup_wait_seconds,
            interval=cfg.startup_retry_interval_seconds)
        if not ready:
            self._ui_q.put(("ollama_status", "unavailable",
                             f"Ollama did not respond in {cfg.startup_wait_seconds} s."))
            return
        self._finish_ollama_check(client, cfg)

    def _finish_ollama_check(self, client, cfg):
        ok, msg = client.check_status()
        if ok:
            self._ui_q.put(("ollama_status", "ready",    cfg.model))
        else:
            self._ui_q.put(("ollama_status", "no_model", cfg.model))

    def _rewrite_selected(self):
        try:
            sel_start = self._corr_box.index(tk.SEL_FIRST)
            sel_end   = self._corr_box.index(tk.SEL_LAST)
            sel_text  = self._corr_box.get(sel_start, sel_end)
        except tk.TclError:
            messagebox.showinfo("No text selected",
                                "Select text in the Corrected panel first.",
                                parent=self)
            return
        if not sel_text.strip():
            messagebox.showinfo("Empty selection",
                                "The selected text is empty.", parent=self)
            return
        if self._rewrite_svc is None:
            messagebox.showerror(
                "Ollama not available",
                f"Install Ollama: https://ollama.ai\n"
                f"Pull model:     ollama pull {self.cfg.llm.model}\n"
                f"Start server:   ollama serve",
                parent=self)
            return

        self._corr_box.config(state="disabled")
        self._rewrite_sel_btn.config(state="disabled")
        self._nb.select(1)
        self._set_status(
            f"Rewriting with {self.cfg.llm.model}…  "
            f"(up to {self.cfg.llm.timeout_seconds} s)", AMBER)

        threading.Thread(
            target=self._rewrite_selected_bg,
            args=(sel_text, sel_start, sel_end),
            daemon=True).start()

    def _rewrite_selected_bg(self, sel_text, sel_start, sel_end):
        try:
            result = self._rewrite_svc.rewrite(sel_text)
            self._ui_q.put(("rewrite_sel_done",
                             sel_text, result, sel_start, sel_end))
        except OllamaConnectionError as exc:
            msg = str(exc)
            if "timed out" in msg.lower() or "timeout" in msg.lower():
                msg = (f"Ollama timed out after {self.cfg.llm.timeout_seconds} s.\n\n"
                       "Try again, or increase timeout_seconds in config/config.yaml.")
            self._ui_q.put(("rewrite_sel_error", msg))
        except OllamaError as exc:
            self._ui_q.put(("rewrite_sel_error", str(exc)))
        except Exception as exc:
            self._ui_q.put(("rewrite_sel_error", f"Unexpected error: {exc}"))
        finally:
            self._ui_q.put(("rewrite_sel_enable", None))

    # ── Font size ─────────────────────────────────────────────────────────────

    def _font_up(self):
        if self._font_size < FONT_SIZE_MAX:
            self._font_size += 1
            self._apply_font()

    def _font_down(self):
        if self._font_size > FONT_SIZE_MIN:
            self._font_size -= 1
            self._apply_font()

    def _apply_font(self):
        self._size_lbl.config(text=str(self._font_size))
        f = ("Segoe UI", self._font_size)
        for box in (self._live_box, self._corr_box):
            box.config(font=f)

    # ── Save / Export ─────────────────────────────────────────────────────────

    def _current_text(self) -> str:
        """Return best available text: corrected > live."""
        t = self._corr_box.get("1.0", "end").strip()
        if not t:
            t = self._live_box.get("1.0", "end").strip()
        return t

    def _save_docx(self):
        text = self._current_text()
        if not text:
            messagebox.showwarning("Nothing to save", "Transcribe something first.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx"), ("All files", "*.*")],
            title="Export as Word document",
            initialfile="dictation.docx")
        if not path:
            return
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = Pt(72)
                sec.left_margin = sec.right_margin  = Pt(90)

            h = doc.add_heading("Pathology Dictation Report", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)

            dp = doc.add_paragraph(
                datetime.datetime.now().strftime("Generated: %d %B %Y  %H:%M"))
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in dp.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x6B, 0x6D, 0x88)

            doc.add_paragraph()
            for block in text.split("\n\n"):
                block = block.strip()
                if block:
                    p = doc.add_paragraph(block)
                    p.paragraph_format.space_after = Pt(8)
                    for run in p.runs:
                        run.font.size = Pt(12)

            doc.save(path)
            self._set_status(f"Exported: {Path(path).name}", GREEN_OK)
        except ImportError:
            messagebox.showerror("python-docx not installed",
                                 "Install with:  pip install python-docx")
        except Exception as e:
            messagebox.showerror("Export failed", str(e))

    # ── Recording ─────────────────────────────────────────────────────────────

    def _toggle_record(self):
        if self.is_loading:
            return
        if not self.transcriber:
            messagebox.showwarning(
                "Not ready",
                "Whisper model is not loaded.\n\n"
                "Use Settings > Whisper Model Folder to select a local model folder,\n"
                "or wait for the model to finish loading.")
            return
        if not self.is_recording:
            self._start_recording()
        else:
            self._stop_and_finalize()

    def _start_recording(self):
        self.is_recording = True
        self._live_text   = ""
        self._write_box(self._live_box, "")
        self._nb.select(0)

        self._rec_btn.config(
            text="⏹   Stop Recording   (F9)",
            bg=REC_HL, activebackground=REC_PR, relief="sunken")
        self._set_status("Recording…  speak clearly", RED_REC)
        self._rec_t0 = time.monotonic()

        self._recorder = AudioRecorder(self.cfg.audio, self.cfg.audio_dir)
        self._recorder.start_recording()
        self._tick()
        threading.Thread(target=self._stream_loop, daemon=True).start()

    def _stop_only(self):
        if not self.is_recording:
            return
        self.is_recording = False
        self._rec_btn.config(
            text="⏺   Start Recording   (F9)",
            bg=REC_BG, activebackground=REC_HL, relief="raised")
        if self._recorder:
            self._recorder.stop_recording()
        self._vu.reset()
        self._set_status("Recording cancelled", TEXT_DIM)
        self._btn_enabled(True)

    def _stop_and_finalize(self):
        self.is_recording = False
        self._vu.reset()
        self._rec_btn.config(
            text="⏺   Start Recording   (F9)",
            bg=REC_BG, activebackground=REC_HL, relief="raised")
        self._btn_enabled(False)
        self._set_status("Transcribing…  please wait", AMBER)
        audio = self._recorder.stop_recording() if self._recorder else None
        sr    = self.cfg.audio.sample_rate
        threading.Thread(
            target=self._finalize, args=(audio, sr), daemon=True).start()

    # ── Streaming ─────────────────────────────────────────────────────────────

    def _stream_loop(self):
        sr           = self.cfg.audio.sample_rate
        need_samples = int(STREAM_CHUNK_SEC * sr)
        done_samples = 0
        while self.is_recording:
            time.sleep(0.25)
            recorder = self._recorder
            if not recorder:
                continue
            try:
                buf = list(recorder.audio_buffer)
            except Exception:
                continue
            total = sum(len(c) for c in buf)
            if total - done_samples < need_samples:
                continue
            if not self._stream_lock.acquire(blocking=False):
                continue
            done_samples = total
            try:
                audio = np.concatenate(buf, axis=0).squeeze().astype(np.float32)
                mx = np.abs(audio).max()
                if mx > 0:
                    audio = audio / mx
                segs, _ = self.transcriber.model.transcribe(
                    audio,
                    language=self.cfg.transcription.language,
                    beam_size=STREAM_BEAM, best_of=1,
                    word_timestamps=False, vad_filter=True, temperature=0.0)
                text = " ".join(s.text for s in segs).strip()
                if text:
                    live_text, _ = self._voice_cmd.process(text)
                    self._live_text = live_text
                    self._ui_q.put(("live", live_text))
            except Exception as e:
                logger.debug(f"stream: {e}")
            finally:
                self._stream_lock.release()

    # ── Final pass ────────────────────────────────────────────────────────────

    def _finalize(self, audio, sr):
        try:
            if audio is None or len(audio) == 0:
                self._ui_q.put(("status",
                                 "No audio captured — check microphone", RED_REC))
                return
            raw = self.transcriber.transcribe(audio, sr)
            if not raw:
                raw = self._live_text
            if not raw:
                self._ui_q.put(("status",
                                 "Nothing detected — please try again", RED_REC))
                return

            raw_after_vc, vc_applied = self._voice_cmd.process(raw)
            corrected, changes = self._corrector.correct_with_logging(raw_after_vc)

            if self.cfg.ui.auto_copy_to_clipboard:
                ClipboardHandler.copy_to_clipboard(corrected)

            self._ui_q.put(("final", raw, corrected, changes, vc_applied))

            n = len(changes)
            v = len(vc_applied)
            parts = [f"Done  ·  {n} correction{'s' if n!=1 else ''} applied"]
            if v:
                parts.append(f"{v} voice command{'s' if v!=1 else ''}")
            if self.cfg.ui.auto_copy_to_clipboard:
                parts.append("copied to clipboard")
            self._ui_q.put(("status", "  ·  ".join(parts), GREEN_OK))

        except Exception as e:
            logger.error(f"finalize: {e}", exc_info=True)
            self._ui_q.put(("status", f"Error: {e}", RED_REC))
        finally:
            self._ui_q.put(("btn", True))

    # ── Timer / VU ────────────────────────────────────────────────────────────

    def _tick(self):
        if not self.is_recording:
            self._time_lbl.config(text="")
            return
        elapsed = time.monotonic() - self._rec_t0
        self._time_lbl.config(text=f"{elapsed:.1f} s")
        try:
            buf = self._recorder.audio_buffer
            if buf:
                recent = np.concatenate(buf[-4:], axis=0)
                rms    = float(np.sqrt(np.mean(recent ** 2)))
                self._vu.update(rms)
        except Exception:
            pass
        self._cursor_on = not self._cursor_on
        cur = "▌" if self._cursor_on else " "
        self._write_live(self._live_text + cur, TEXT_LIVE)
        self.after(200, self._tick)

    # ── Queue poll ────────────────────────────────────────────────────────────

    def _poll(self):
        try:
            while True:
                m  = self._ui_q.get_nowait()
                op = m[0]

                if op == "status":
                    self._set_status(m[1], m[2])

                elif op == "btn":
                    self._btn_enabled(m[1])

                elif op == "model_status":
                    _, status, _path = m
                    self._update_model_label(status)

                elif op == "live":
                    if self.is_recording:
                        cur = "▌" if self._cursor_on else " "
                        self._write_live(m[1] + cur, TEXT_LIVE)

                elif op == "final":
                    _, raw, corrected, changes, vc_applied = m
                    self._write_box(self._live_box, corrected, TEXT)
                    self._insert_at_cursor_in_corr(corrected)
                    self._nb.select(1)

                    parts = []
                    if vc_applied:
                        parts.append("  Voice cmds: " +
                                     ", ".join(f"[{c}]" for c in vc_applied))
                    if changes:
                        parts.append(
                            "  Corrections: " +
                            "   ·   ".join(
                                f"'{c['original']}' → '{c['replacement']}'"
                                for c in changes))
                    self._changes_lbl.config(
                        text="   ".join(parts) if parts
                        else "  No voice commands or terminology corrections applied.")

                elif op == "rewrite_sel_done":
                    _, original, rewritten, sel_start, sel_end = m
                    self._corr_box.config(state="normal")
                    self._set_status(
                        "Rewrite ready — review in preview dialog", GREEN_OK)
                    def _on_accept(final_text, _ss=sel_start, _se=sel_end):
                        self._corr_box.delete(_ss, _se)
                        self._corr_box.insert(_ss, final_text)
                        self._corr_box.edit_separator()
                        self._mark_dirty()
                        self._set_status("Rewrite accepted", GREEN_OK)
                    def _on_reject():
                        self._set_status("Rewrite rejected — original kept", TEXT_DIM)
                    RewritePreviewDialog(
                        self, original, rewritten, _on_accept, _on_reject)

                elif op == "rewrite_sel_error":
                    _, msg = m
                    self._corr_box.config(state="normal")
                    self._set_status("Rewrite error", RED_REC)
                    messagebox.showerror("Rewrite failed", msg, parent=self)

                elif op == "rewrite_sel_enable":
                    self._rewrite_sel_btn.config(state="normal")

                elif op == "ollama_foot":
                    _, text, color = m
                    self._foot.config(text=text, fg=color)

                elif op == "ollama_status":
                    _, state, detail = m
                    if state == "ready":
                        self._rewrite_sel_btn.config(state="normal")
                        if hasattr(self, "_rewrite_menu"):
                            self._rewrite_menu.entryconfig(0, state="normal")
                        self._foot.config(
                            text=f"Ollama: {detail}  ✓", fg=GREEN_OK)
                    elif state == "no_model":
                        self._rewrite_sel_btn.config(state="disabled")
                        if hasattr(self, "_rewrite_menu"):
                            self._rewrite_menu.entryconfig(0, state="disabled")
                        self._foot.config(
                            text=f"Ollama: '{detail}' not installed  —  "
                                 f"run: ollama pull {detail}",
                            fg=AMBER)
                        messagebox.showwarning(
                            "Qwen model not installed",
                            f"Ollama is running but '{detail}' is not installed.\n\n"
                            f"Pull it with:\n  ollama pull {detail}\n\n"
                            "Dictation works normally.",
                            parent=self)
                    elif state == "unavailable":
                        self._rewrite_sel_btn.config(state="disabled")
                        if hasattr(self, "_rewrite_menu"):
                            self._rewrite_menu.entryconfig(0, state="disabled")
                        self._foot.config(
                            text="Ollama unavailable — rewrite disabled  "
                                 "(dictation works normally)",
                            fg=AMBER)

                elif op == "msgbox":
                    _, lvl, title, txt = m
                    (messagebox.showerror if lvl == "error"
                     else messagebox.showinfo)(title, txt)

        except queue.Empty:
            pass
        self.after(40, self._poll)

    # ── Widget helpers ────────────────────────────────────────────────────────

    def _write_live(self, text: str, color=TEXT_LIVE):
        box = self._live_box
        box.config(state="normal")
        box.delete("1.0", "end")
        if text:
            box.insert("end", text)
            box.tag_add("c", "1.0", "end")
            box.tag_config("c", foreground=color)
        box.see("end")
        box.config(state="disabled")

    def _write_box(self, box, text: str, color=TEXT, stay_enabled=False):
        box.config(state="normal")
        box.delete("1.0", "end")
        if text:
            box.insert("end", text)
            box.tag_add("c", "1.0", "end")
            box.tag_config("c", foreground=color)
        if not stay_enabled:
            box.config(state="disabled")

    def _set_status(self, text: str, color=TEXT_DIM):
        self._stat_lbl.config(text=text, fg=color)
        self._foot.config(text=text, fg=color)

    def _btn_enabled(self, on: bool):
        self._rec_btn.config(
            state="normal" if on else "disabled",
            cursor="hand2" if on else "arrow")

    def _copy_result(self):
        text = self._current_text()
        if text:
            ClipboardHandler.copy_to_clipboard(text)
            self._set_status("Copied to clipboard!", GREEN_OK)
        else:
            self._set_status("Nothing to copy yet.", TEXT_DIM)

    def _clear_all(self):
        if self._doc_dirty:
            if not messagebox.askyesno(
                    "Unsaved changes",
                    "Clear all text?  Unsaved changes will be lost.",
                    parent=self):
                return
        self._loading_file = True
        self._write_box(self._live_box, "")
        self._corr_box.delete("1.0", "end")
        self._corr_box.edit_modified(False)
        self._changes_lbl.config(text="")
        self._vu.reset()
        self._time_lbl.config(text="")
        self._live_text = ""
        self._doc_path  = None
        self._doc_dirty = False
        self._loading_file = False
        self._update_title()

    # ── Close ─────────────────────────────────────────────────────────────────

    # ── AI rewrite ────────────────────────────────────────────────────────────

    def _populate_rewrite_models(self):
        """Scan models/rewrite/ for .gguf files and populate the combobox."""
        try:
            models = scan_models(self.cfg.models_dir)
        except Exception:
            models = []
        if models:
            names = [m.name for m in models]
            self._rewr_model_cb.configure(values=names)
            # Keep current selection if still valid, else pick first
            if self._rewr_model_var.get() not in names:
                self._rewr_model_var.set(names[0])
            self._rewrite_btn.config(state="normal")
        else:
            self._rewr_model_cb.configure(values=["No models found"])
            self._rewr_model_var.set("No models found")
            self._rewrite_btn.config(state="disabled")

    def _rewrite(self):
        """Start AI rewrite of the Corrected text in a background thread."""
        text = self._corr_box.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning(
                "Nothing to rewrite",
                "Dictate and correct something first.",
                parent=self)
            return

        model_name = self._rewr_model_var.get()
        if not model_name or model_name == "No models found":
            messagebox.showinfo(
                "No rewrite model found",
                "Place a .gguf model file in:\n"
                f"  {self.cfg.models_dir / 'rewrite'}\n\n"
                "Supported models (Q4_K_M GGUF format):\n"
                "  • Qwen2.5-1.5B-Instruct-Q4_K_M.gguf  (~0.9 GB)\n"
                "  • Qwen2.5-3B-Instruct-Q4_K_M.gguf    (~1.8 GB)\n"
                "  • Phi-3.5-mini-instruct-Q4_K_M.gguf  (~2.2 GB)\n"
                "  • Llama-3.2-1B-Instruct-Q4_K_M.gguf  (~0.7 GB)\n\n"
                "Download from HuggingFace — see README.",
                parent=self)
            return

        # Check llama-cpp-python is available before launching thread
        avail, msg = LocalRewriter.check_available()
        if not avail:
            messagebox.showerror("llama-cpp-python not installed", msg, parent=self)
            return

        model_path = self.cfg.models_dir / "rewrite" / model_name
        self._rewrite_btn.config(state="disabled")
        self._set_status(
            f"Loading  {model_name}…  (first run may take 10–30 s)", AMBER)
        threading.Thread(
            target=self._rewrite_bg, args=(text, model_path), daemon=True
        ).start()

    def _rewrite_bg(self, text: str, model_path: Path):
        """Background worker: load GGUF model if needed, then rewrite."""
        try:
            # Load or swap model (unload old if different)
            if self._rewriter is None or self._rewriter.model_path != model_path:
                if self._rewriter and self._rewriter.is_loaded():
                    self._rewriter.unload()
                self._rewriter = LocalRewriter(model_path)
                self._rewriter.load()

            self._ui_q.put(("status", "Rewriting text…", AMBER))
            result = self._rewriter.rewrite(text)
            self._ui_q.put(("rewrite_done", result))
            self._ui_q.put(("status",
                             "Rewrite complete  ·  appended to  ✏ Corrected  tab", GREEN_OK))
        except Exception as e:
            logger.error(f"rewrite_bg: {e}", exc_info=True)
            self._ui_q.put(("status", f"Rewrite error: {e}", RED_REC))
            self._ui_q.put(("msgbox", "error", "Rewrite failed", str(e)))
        finally:
            self._ui_q.put(("rewrite_btn_enable", True))

    def _on_close(self):
        if self._doc_dirty:
            ans = messagebox.askyesnocancel(
                "Unsaved changes",
                "You have unsaved changes.\nSave before exiting?",
                parent=self)
            if ans is None:
                return
            if ans:
                if not self._save_file():
                    return
        if self._autosave_id:
            self.after_cancel(self._autosave_id)
        self.cfg.save_user_settings()
        self.destroy()


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()
